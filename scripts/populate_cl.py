"""
populate_cl.py — Fill the cover letter template for a specific job application.

Reads cl_text/cv_text from Supabase (--db-id mode) or from a Notion page (legacy),
maps CL paragraphs to placeholders, and saves a finished .docx.

Usage:
    py scripts/populate_cl.py --db-id <job_id> [--lang fr|en]
    py scripts/populate_cl.py <notion_page_id> [--lang fr|en]    (legacy)
    py scripts/populate_cl.py <notion_page_id> "Company_JobTitle" [--lang fr|en]

Templates:
    FR (default): cl_template.docx     (French CL layout — Raydiall base)
    EN:           cl_template_en.docx  (English CL layout)

DB mode: reads PG_CONN from env var or supabase_connection_string in config.json.
Notion token: read from .mcp.json in the repo root, or set NOTION_API_TOKEN env var.
"""

import sys
import os
import json
import re
import subprocess
from datetime import date
from pathlib import Path
from docx import Document
from notion_client import Client

REPO_ROOT = Path(__file__).parent.parent
TEMPLATES_DIR = REPO_ROOT / "templates"
OUTPUT_DIR = REPO_ROOT / "outputs"

TEMPLATES = {
    "fr": str(TEMPLATES_DIR / "cl_template.docx"),
    "en": str(TEMPLATES_DIR / "cl_template_en.docx"),
}

MONTHS_FR = {
    1: "janvier", 2: "février", 3: "mars", 4: "avril",
    5: "mai", 6: "juin", 7: "juillet", 8: "août",
    9: "septembre", 10: "octobre", 11: "novembre", 12: "décembre",
}


def get_notion_token():
    token = os.environ.get("NOTION_API_TOKEN")
    if token:
        return token
    mcp_path = REPO_ROOT / ".mcp.json"
    if mcp_path.exists():
        with open(mcp_path) as f:
            data = json.load(f)
        try:
            return data["mcpServers"]["notion"]["env"]["NOTION_API_TOKEN"]
        except KeyError:
            pass
    print("ERROR: No Notion token found.")
    print("  Set NOTION_API_TOKEN env var, or ensure .mcp.json exists with the token.")
    sys.exit(1)


def get_pg_conn():
    """Return (pg_conn_string, pg_module_path) from env or config.json."""
    conn = os.environ.get("PG_CONN")
    module = os.environ.get("PG_MODULE", "pg")
    if conn:
        return conn, module
    config_path = REPO_ROOT / "config.json"
    if config_path.exists():
        with open(config_path) as f:
            cfg = json.load(f)
        return cfg.get("supabase_connection_string", ""), cfg.get("pg_module_path", "pg")
    print("ERROR: No PG_CONN env var and no config.json found.")
    sys.exit(1)


def fetch_job_from_db(job_id):
    """Fetch cl_text, cv_text, company, job_title, location from job_applications."""
    pg_conn, pg_module = get_pg_conn()
    script = f"""
const {{Client}} = require('{pg_module}');
const c = new Client({{connectionString: '{pg_conn}'}});
c.connect()
  .then(() => c.query(
    'SELECT cl_text, cv_text, company, job_title, location FROM job_applications WHERE id = $1',
    [{job_id}]
  ))
  .then(r => {{ console.log(JSON.stringify(r.rows)); return c.end(); }})
  .catch(e => {{ console.error(e.message); process.exit(1); }});
"""
    result = subprocess.run(["node", "-e", script], capture_output=True, text=True, encoding="utf-8")
    if result.returncode != 0:
        print(f"ERROR querying DB: {result.stderr.strip()}")
        sys.exit(1)
    rows = json.loads(result.stdout.strip())
    if not rows:
        print(f"ERROR: No row found in job_applications with id={job_id}")
        sys.exit(1)
    return rows[0]


def fetch_page_blocks(notion, page_id):
    blocks = []
    cursor = None
    while True:
        response = notion.blocks.children.list(block_id=page_id, start_cursor=cursor)
        blocks.extend(response["results"])
        if not response["has_more"]:
            break
        cursor = response["next_cursor"]
    return blocks


def block_text(block):
    btype = block["type"]
    rich_text = block.get(btype, {}).get("rich_text", [])
    return "".join(rt.get("plain_text", "") for rt in rich_text)


def extract_section(blocks, section_heading):
    """Extract lines from a ## section until the next ## section."""
    in_section = False
    lines = []
    for block in blocks:
        btype = block["type"]
        text = block_text(block).strip()
        if btype == "heading_2" and section_heading in text:
            in_section = True
            continue
        if in_section:
            if btype == "heading_2":
                break
            if text:
                lines.append(text)
    return lines


def parse_application_notes(lines):
    """
    Extract company, job title, and location from Application Notes section.
    Looks for lines like:
      **Role:** Job Title @ Company
      **Location/Zone:** Voiron (38)
    """
    company = ""
    job_title = ""
    location = ""
    for line in lines:
        if "**Role:**" in line or "Role:" in line:
            # Strip bold markers or plain "Role:" prefix, then extract title @ company
            clean = re.sub(r"\*{0,2}Role:\*{0,2}\s*", "", line).strip()
            if "@" in clean:
                parts = clean.split("@", 1)
                job_title = parts[0].strip()
                company = parts[1].strip()
        if "**Location" in line or "Location:" in line:
            clean = re.sub(r"\*{0,2}Location[^:]*:\*{0,2}\s*", "", line).strip()
            # Remove zone emoji if present
            clean = re.sub(r"[🟢🟡🟠🔴🌐]", "", clean).strip()
            location = clean.split("·")[0].strip()  # take city part before any ·
    return company, job_title, location


def get_base_city():
    """Read base city from .mcp.json (written by setup.py from .env profile)."""
    mcp_path = REPO_ROOT / ".mcp.json"
    if mcp_path.exists():
        try:
            with open(mcp_path) as f:
                data = json.load(f)
            return data.get("base_city", "")
        except Exception:
            pass
    return ""


def french_date():
    today = date.today()
    city = get_base_city()
    prefix = f"{city}, le " if city else ""
    return f"{prefix}{today.day} {MONTHS_FR[today.month]} {today.year}"


def replace_placeholder(doc, placeholder, new_text):
    """Replace a placeholder anywhere in the document (paragraphs + tables)."""
    replaced = False

    # Top-level paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, new_text)
                replaced = True

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, new_text)
                            replaced = True

    return replaced


def delete_placeholder_paragraph(doc, placeholder):
    """Remove the paragraph containing a placeholder entirely (used when value is empty)."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            p = para._element
            p.getparent().remove(p)
            return True
    return False


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    # Parse args: [--db-id <id> | <notion_page_id>] [custom_name] [--lang fr|en]
    args = sys.argv[1:]
    lang = "fr"
    if "--lang" in args:
        idx = args.index("--lang")
        lang = args[idx + 1].lower() if idx + 1 < len(args) else "fr"
        args = [a for i, a in enumerate(args) if i != idx and i != idx + 1]

    if lang not in TEMPLATES:
        print(f"ERROR: Unknown language '{lang}'. Use 'fr' or 'en'.")
        sys.exit(1)

    template_path = TEMPLATES[lang]

    # DB mode: --db-id <job_id>
    if args[0] == "--db-id":
        if len(args) < 2:
            print("ERROR: --db-id requires a job id argument.")
            sys.exit(1)
        job_id = args[1]
        custom_name = args[2] if len(args) >= 3 else None

        print(f"Fetching cl_text/cv_text from DB for job id={job_id}...")
        row = fetch_job_from_db(job_id)

        cl_text = row.get("cl_text") or ""
        cv_text = row.get("cv_text") or ""
        company = row.get("company") or ""
        job_title = row.get("job_title") or ""
        raw_location = row.get("location") or ""

        if not cl_text.strip():
            print(f"ERROR: cl_text is empty for job id={job_id}. Run /job-apply first.")
            sys.exit(1)

        page_title = f"{company} — {job_title}"

        # Parse CL paragraphs from plain text (blank-line separated)
        cl_paras = [p.strip() for p in re.split(r"\n\s*\n", cl_text) if p.strip() and len(p.strip()) > 30]

        # CV headline: first line with '|' (skips candidate name line)
        cv_lines = [l.strip() for l in cv_text.splitlines() if l.strip()]
        cl_headline = next((l for l in cv_lines if "|" in l), cv_lines[0] if cv_lines else job_title)

        # Location: use raw location from DB; add ", France" if needed
        location = raw_location.strip()
        if location and not location.lower().endswith("france"):
            location = location.rstrip(", ") + ", France"

    else:
        # Legacy Notion mode
        page_id = args[0].replace("-", "")
        custom_name = args[1] if len(args) >= 2 else None

        token = get_notion_token()
        notion = Client(auth=token)

        print(f"Fetching Notion page {page_id}...")
        try:
            page_meta = notion.pages.retrieve(page_id=page_id)
        except Exception as e:
            print(f"ERROR fetching page: {e}")
            sys.exit(1)

        title_prop = page_meta.get("properties", {}).get("title", {}).get("title", [])
        page_title = "".join(t.get("plain_text", "") for t in title_prop)
        print(f"Page title: {page_title}")

        blocks = fetch_page_blocks(notion, page_id)

        cl_lines = extract_section(blocks, "Cover Letter")
        if not cl_lines:
            print("ERROR: 'Cover Letter' section not found in the Notion page.")
            sys.exit(1)

        notes_lines = extract_section(blocks, "Application Notes")
        company, job_title, location = parse_application_notes(notes_lines)
        if location and not location.lower().endswith("france"):
            location = location.rstrip(", ") + ", France"

        cv_lines = extract_section(blocks, "Tailored CV")
        cl_headline = next((l for l in cv_lines if "|" in l), cv_lines[0] if cv_lines else job_title)

        cl_paras = [l for l in cl_lines if len(l) > 30]

    # Map paragraphs to placeholders
    # Expected order: opening, body1, body2, body3 (optional), closing
    opening = cl_paras[0] if len(cl_paras) > 0 else ""
    body1   = cl_paras[1] if len(cl_paras) > 1 else ""
    body2   = cl_paras[2] if len(cl_paras) > 2 else ""
    body3   = cl_paras[3] if len(cl_paras) > 3 else ""
    closing = cl_paras[4] if len(cl_paras) > 4 else (cl_paras[3] if len(cl_paras) > 3 else "")

    if len(cl_paras) == 4:
        body3   = ""
        closing = cl_paras[3]

    print(f"\nExtracted:")
    print(f"  Company    : {company}")
    print(f"  Job title  : {job_title}")
    print(f"  Location   : {location}")
    print(f"  CL headline: {cl_headline}")
    print(f"  Paragraphs : {len(cl_paras)}")
    print(f"  Opening    : {opening[:60]}...")

    print(f"  Language   : {lang.upper()} -> {Path(template_path).name}")

    if not os.path.exists(template_path):
        script = "make_cl_template.py" if lang == "fr" else "make_cl_template_en.py"
        print(f"\nERROR: Template not found: {template_path}")
        print(f"Run 'py scripts/{script}' first.")
        sys.exit(1)

    doc = Document(template_path)

    replacements = {
        "{{CL_HEADLINE}}":       cl_headline,
        "{{DATE}}":              french_date(),
        "{{COMPANY_ADDRESSEE}}": f"{company} — Service Recrutement" if company else "{{COMPANY_ADDRESSEE}}",
        "{{COMPANY_LOCATION}}":  location,
        "{{SUBJECT_LINE}}":      f"Objet : Candidature — {job_title}" if job_title else "{{SUBJECT_LINE}}",
        "{{OPENING_PARA}}":      opening,
        "{{BODY_PARA_1}}":       body1,
        "{{BODY_PARA_2}}":       body2,
        "{{BODY_PARA_3}}":       body3,
        "{{CLOSING_PARA}}":      closing,
    }

    for placeholder, value in replacements.items():
        if not value and placeholder == "{{BODY_PARA_3}}":
            # Delete the paragraph entirely rather than leaving a blank line
            ok = delete_placeholder_paragraph(doc, placeholder)
            status = "[DELETED — empty]" if ok else "[NOT FOUND]"
        else:
            ok = replace_placeholder(doc, placeholder, value)
            status = "[OK]" if ok else "[NOT FOUND]"
        print(f"  {placeholder}: {status}")

    # Output filename
    if custom_name:
        filename = custom_name if custom_name.endswith(".docx") else custom_name + ".docx"
    else:
        safe_title = re.sub(r'[<>:"/\\|?*]', "", page_title).strip()
        filename = (safe_title + "_CL.docx") if safe_title else "CL_output.docx"

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    print(f"\nDone. CL saved to:\n  {output_path}")


if __name__ == "__main__":
    main()
