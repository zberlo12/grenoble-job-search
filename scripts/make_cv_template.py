"""
make_cv_template.py — Run ONCE per CV variant to create a Word template.

Reads your base resume .docx, replaces the job title line and profile summary
with {{CV_HEADLINE}} and {{PROFILE_SUMMARY}} placeholders, and saves a
template file to templates/. The populate_cv.py script fills those placeholders
per application.

Usage:
    py scripts/make_cv_template.py <source_cv.docx> <output_name>

    <source_cv.docx>  Path to your existing formatted CV Word file
    <output_name>     Template name without .docx (e.g. cv_template_fpa_fr)

Example:
    py scripts/make_cv_template.py "C:/My Documents/My_CV.docx" cv_template_fpa_fr

The script looks for your full name (in uppercase) to locate the headline line.
Your full name is read from .mcp.json (set via setup.py from your .env profile).
"""

import sys
import json
from pathlib import Path
from docx import Document

REPO_ROOT = Path(__file__).parent.parent
TEMPLATES_DIR = REPO_ROOT / "templates"


def get_full_name_upper():
    """Read full name from .mcp.json, return in uppercase for matching."""
    mcp_path = REPO_ROOT / ".mcp.json"
    if mcp_path.exists():
        try:
            with open(mcp_path) as f:
                data = json.load(f)
            name = data.get("full_name", "")
            return name.upper() if name else ""
        except Exception:
            pass
    return ""


def replace_paragraph_text(para, new_text):
    """Replace a paragraph's text while preserving the formatting of the first run."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    source_path = sys.argv[1]
    output_name = sys.argv[2]
    if not output_name.endswith(".docx"):
        output_name += ".docx"

    full_name_upper = get_full_name_upper()
    if not full_name_upper:
        print("ERROR: Could not read full_name from .mcp.json.")
        print("  Run setup.py --profile <name> first, and ensure FULL_NAME is in your .env file.")
        sys.exit(1)

    if not Path(source_path).exists():
        print(f"ERROR: Source file not found: {source_path}")
        sys.exit(1)

    doc = Document(source_path)

    table = doc.tables[0]
    right_cell = table.rows[0].cells[1]
    paras = right_cell.paragraphs

    after_name = False
    headline_replaced = False
    after_profile = False
    summary_replaced = False

    for para in paras:
        text = para.text.strip()
        if not text:
            continue

        if text == full_name_upper:
            after_name = True
            continue

        if after_name and not headline_replaced:
            print(f"Replacing headline: '{text}'")
            replace_paragraph_text(para, "{{CV_HEADLINE}}")
            headline_replaced = True
            after_name = False
            continue

        if text == "PROFILE":
            after_profile = True
            continue

        if after_profile and not summary_replaced:
            print(f"Replacing profile: '{text[:80]}...'")
            replace_paragraph_text(para, "{{PROFILE_SUMMARY}}")
            summary_replaced = True
            after_profile = False
            break

    if not headline_replaced:
        print(f"ERROR: Could not find the headline paragraph (looked for name: '{full_name_upper}').")
        return
    if not summary_replaced:
        print("ERROR: Could not find the PROFILE section.")
        return

    TEMPLATES_DIR.mkdir(exist_ok=True)
    output_path = TEMPLATES_DIR / output_name
    doc.save(str(output_path))
    print(f"\nTemplate saved: {output_path}")
    print("Run populate_cv.py to fill it for a specific application.")


if __name__ == "__main__":
    main()
