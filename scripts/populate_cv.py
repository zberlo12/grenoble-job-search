"""
populate_cv.py — Fill the CV template for a specific job application.

Reads cv_text from Supabase (--db-id mode) or from a Notion page (legacy),
extracts the headline and profile summary, fills them into the CV template,
and saves a finished .docx ready for formatting review.

Usage:
    py scripts/populate_cv.py --db-id <job_id> [--approach fpa-fr]
    py scripts/populate_cv.py <notion_page_id> [--approach fpa-fr]    (legacy)
    py scripts/populate_cv.py <notion_page_id> "Company_JobTitle" [--approach fpa-fr]

Approach -> template mapping:
    fpa-fr         CV — FP&A Focus — FR     (Hays CDG/FPA base)
    costcontrol-fr CV — Cost Control — FR   (Teledyne CDG base)
    raf-fr         CV — Standard — FR       (Raydiall RAF base — callback-proven)
    fpa-en         CV — FP&A Focus — EN     (ThermoFisher Sr Mgr base)
    hof-en         CV — Head of Finance EN  (ESRF HoF base — Director-level EN)

DB mode: reads PG_CONN from env var or supabase_connection_string in config.json.
Notion token: read from .mcp.json in the repo root, or set NOTION_API_TOKEN env var.
"""

import sys
import os
import json
import re
import subprocess
from pathlib import Path
from docx import Document
from notion_client import Client

REPO_ROOT = Path(__file__).parent.parent
TEMPLATES_DIR = REPO_ROOT / "templates"
OUTPUT_DIR = REPO_ROOT / "outputs"

TEMPLATES = {
    "fpa-fr":         str(TEMPLATES_DIR / "cv_template_fpa_fr.docx"),
    "costcontrol-fr": str(TEMPLATES_DIR / "cv_template_costcontrol_fr.docx"),
    "raf-fr":         str(TEMPLATES_DIR / "cv_template_raf_fr.docx"),
    "fpa-en":         str(TEMPLATES_DIR / "cv_template_fpa_en.docx"),
    "hof-en":         str(TEMPLATES_DIR / "cv_template_hof_en.docx"),
    # legacy aliases
    "fr":             str(TEMPLATES_DIR / "cv_template_fpa_fr.docx"),
    "en":             str(TEMPLATES_DIR / "cv_template_fpa_en.docx"),
}


def get_notion_token():
    """Read Notion token from NOTION_API_TOKEN env var or .mcp.json."""
    token = os.environ.get("NOTION_API_TOKEN")
    if token:
        return token
    mcp_path = REPO_ROOT / ".mcp.json"
    if mcp_path.exists():
        with open(mcp_path) as f:
            data = json.load(f)
        try:
            return data["mcpServers"]["notion"]["env"]["NOTION_API_TOKEN"]
        except KeyError:
            pass
    print("ERROR: No Notion token found.")
    print("  Set NOTION_API_TOKEN env var, or ensure .mcp.json exists with the token.")
    sys.exit(1)


def get_pg_conn():
    """Return (pg_conn_string, pg_module_path) from env or config.json."""
    conn = os.environ.get("PG_CONN")
    module = os.environ.get("PG_MODULE", "pg")
    if conn:
        return conn, module
    config_path = REPO_ROOT / "config.json"
    if config_path.exists():
        with open(config_path) as f:
            cfg = json.load(f)
        return cfg.get("supabase_connection_string", ""), cfg.get("pg_module_path", "pg")
    print("ERROR: No PG_CONN env var and no config.json found.")
    sys.exit(1)


def fetch_cv_from_db(job_id):
    """Fetch cv_text and job_title from job_applications by id."""
    pg_conn, pg_module = get_pg_conn()
    script = f"""
const {{Client}} = require('{pg_module}');
const c = new Client({{connectionString: '{pg_conn}'}});
c.connect()
  .then(() => c.query('SELECT cv_text, job_title, company FROM job_applications WHERE id = $1', [{job_id}]))
  .then(r => {{ console.log(JSON.stringify(r.rows)); return c.end(); }})
  .catch(e => {{ console.error(e.message); process.exit(1); }});
"""
    result = subprocess.run(["node", "-e", script], capture_output=True, text=True)
    if result.returncode != 0:
        print(f"ERROR querying DB: {result.stderr.strip()}")
        sys.exit(1)
    rows = json.loads(result.stdout.strip())
    if not rows:
        print(f"ERROR: No row found in job_applications with id={job_id}")
        sys.exit(1)
    return rows[0]


def fetch_page_blocks(notion, page_id):
    """Fetch all blocks from a Notion page, handling pagination."""
    blocks = []
    cursor = None
    while True:
        response = notion.blocks.children.list(block_id=page_id, start_cursor=cursor)
        blocks.extend(response["results"])
        if not response["has_more"]:
            break
        cursor = response["next_cursor"]
    return blocks


def block_text(block):
    """Extract plain text from a Notion block."""
    btype = block["type"]
    rich_text = block.get(btype, {}).get("rich_text", [])
    return "".join(rt.get("plain_text", "") for rt in rich_text)


def extract_cv_section(blocks):
    """
    Extract lines from the 'Tailored CV' section of the Notion page.
    Returns a list of non-empty text strings until the next h2 section.
    """
    in_cv = False
    lines = []

    for block in blocks:
        btype = block["type"]
        text = block_text(block).strip()

        if btype == "heading_2" and "Tailored CV" in text:
            in_cv = True
            continue

        if in_cv:
            if btype == "heading_2":
                break  # Hit the next section (Cover Letter)
            if text:
                lines.append(text)

    return lines


def parse_headline_and_summary(cv_lines):
    """
    Parse the CV headline and profile summary from the CV section lines.

    Expected structure (as drafted by /job-apply):
      Line 0: CV headline (short — the tailored job title variant)
      Later:  Profile summary (the longest paragraph-style line)
    """
    if not cv_lines:
        return None, None

    headline = cv_lines[0]

    # Profile summary: first line over 100 chars that isn't a section header or bullet
    summary = None
    for line in cv_lines[1:]:
        if len(line) > 100 and not line.startswith("##") and not line.startswith("-"):
            summary = line
            break

    # Fallback: second line if no long paragraph found
    if not summary and len(cv_lines) > 1:
        summary = cv_lines[1]

    return headline, summary


def replace_placeholder(cell, placeholder, new_text):
    """Replace a {{placeholder}} in a table cell, preserving run formatting."""
    for para in cell.paragraphs:
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, new_text)
                return True
    return False


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    # Parse args: [--db-id <id> | <notion_page_id>] [custom_name] [--approach <key>]
    args = sys.argv[1:]
    approach = "fpa-fr"
    for flag in ("--approach", "--lang"):
        if flag in args:
            idx = args.index(flag)
            approach = args[idx + 1].lower() if idx + 1 < len(args) else "fpa-fr"
            args = [a for i, a in enumerate(args) if i != idx and i != idx + 1]
            break

    if approach not in TEMPLATES:
        valid = ", ".join(k for k in TEMPLATES if "-" in k)
        print(f"ERROR: Unknown approach '{approach}'. Valid: {valid}")
        sys.exit(1)

    template_path = TEMPLATES[approach]

    # DB mode: --db-id <job_id>
    if args[0] == "--db-id":
        if len(args) < 2:
            print("ERROR: --db-id requires a job id argument.")
            sys.exit(1)
        job_id = args[1]
        custom_name = args[2] if len(args) >= 3 else None

        print(f"Fetching cv_text from DB for job id={job_id}...")
        row = fetch_cv_from_db(job_id)
        cv_text = row.get("cv_text") or ""
        page_title = f"{row.get('company', '')} — {row.get('job_title', '')}"

        if not cv_text.strip():
            print(f"ERROR: cv_text is empty for job id={job_id}. Run /job-apply first.")
            sys.exit(1)

        cv_lines = [l.strip() for l in cv_text.splitlines() if l.strip()]

    else:
        # Legacy Notion mode
        page_id = args[0].replace("-", "")
        custom_name = args[1] if len(args) >= 2 else None

        token = get_notion_token()
        notion = Client(auth=token)

        print(f"Fetching Notion page {page_id}...")
        try:
            page_meta = notion.pages.retrieve(page_id=page_id)
        except Exception as e:
            print(f"ERROR fetching page: {e}")
            sys.exit(1)

        title_prop = page_meta.get("properties", {}).get("title", {}).get("title", [])
        page_title = "".join(t.get("plain_text", "") for t in title_prop)
        if not page_title:
            page_title = page_meta.get("child_page", {}).get("title", "CV_output")

        print(f"Page title: {page_title}")

        blocks = fetch_page_blocks(notion, page_id)
        cv_lines = extract_cv_section(blocks)

        if not cv_lines:
            print("ERROR: 'Tailored CV' section not found in the Notion page.")
            print("       Make sure this is an Application Document page created by /job-apply.")
            sys.exit(1)

    headline, summary = parse_headline_and_summary(cv_lines)

    print(f"\nExtracted:")
    print(f"  Headline : {headline}")
    print(f"  Summary  : {summary[:80]}..." if summary and len(summary) > 80 else f"  Summary  : {summary}")

    print(f"  Approach : {approach} -> {Path(template_path).name}")

    if not os.path.exists(template_path):
        script = "make_cv_template.py"
        print(f"\nERROR: Template not found:\n  {template_path}")
        print(f"Run 'py scripts/{script}' first to create it.")
        sys.exit(1)

    doc = Document(template_path)
    table = doc.tables[0]
    right_cell = table.rows[0].cells[1]

    h_ok = replace_placeholder(right_cell, "{{CV_HEADLINE}}", headline or "")
    s_ok = replace_placeholder(right_cell, "{{PROFILE_SUMMARY}}", summary or "")

    if not h_ok:
        print("WARNING: {{CV_HEADLINE}} placeholder not found in template.")
    if not s_ok:
        print("WARNING: {{PROFILE_SUMMARY}} placeholder not found in template.")

    # Build output filename
    if custom_name:
        filename = custom_name if custom_name.endswith(".docx") else custom_name + ".docx"
    else:
        safe_title = re.sub(r'[<>:"/\\|?*]', "", page_title).strip()
        filename = safe_title + ".docx" if safe_title else "CV_output.docx"

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)

    print(f"\nDone. CV saved to:\n  {output_path}")


if __name__ == "__main__":
    main()
