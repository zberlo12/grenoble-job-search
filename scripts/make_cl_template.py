"""
make_cl_template.py — Run ONCE to create the cover letter Word template.

Reads your existing cover letter .docx and replaces all variable content with
{{placeholders}}, saving cl_template.docx to templates/. The populate_cl.py
script fills those placeholders per application.

Usage:
    py scripts/make_cl_template.py <source_cl.docx> [--lang fr|en]

    <source_cl.docx>  Path to your existing cover letter Word file
    --lang            Language (fr or en, default: fr)

Example:
    py scripts/make_cl_template.py "C:/My Documents/My_CL_French.docx" --lang fr

Your full name is read from .mcp.json (set via setup.py from your .env profile).
"""

import sys
import json
from pathlib import Path
from docx import Document

REPO_ROOT = Path(__file__).parent.parent
TEMPLATES_DIR = REPO_ROOT / "templates"


def get_full_name():
    """Read full name from .mcp.json."""
    mcp_path = REPO_ROOT / ".mcp.json"
    if mcp_path.exists():
        try:
            with open(mcp_path) as f:
                data = json.load(f)
            return data.get("full_name", "")
        except Exception:
            pass
    return ""


def replace_para_text(para, new_text):
    """Replace paragraph text, preserving first run's formatting."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def replace_cell_text(cell, old_substring, new_text):
    """Replace a paragraph in a table cell that contains old_substring."""
    for para in cell.paragraphs:
        if old_substring in para.text:
            replace_para_text(para, new_text)
            return True
    return False


def main():
    args = sys.argv[1:]
    lang = "fr"
    if "--lang" in args:
        idx = args.index("--lang")
        lang = args[idx + 1].lower() if idx + 1 < len(args) else "fr"
        args = [a for i, a in enumerate(args) if i != idx and i != idx + 1]

    if not args:
        print(__doc__)
        sys.exit(1)

    source_path = args[0]
    output_name = "cl_template.docx" if lang == "fr" else "cl_template_en.docx"

    full_name = get_full_name()
    if not full_name:
        print("ERROR: Could not read full_name from .mcp.json.")
        print("  Run setup.py --profile <name> first, and ensure FULL_NAME is in your .env file.")
        sys.exit(1)

    if not Path(source_path).exists():
        print(f"ERROR: Source file not found: {source_path}")
        sys.exit(1)

    doc = Document(source_path)

    # Header table: replace job title line with placeholder
    if doc.tables:
        header_table = doc.tables[0]
        left_cell = header_table.rows[0].cells[0]
        # Find and replace any non-name paragraph in the header (the job title line)
        for para in left_cell.paragraphs:
            text = para.text.strip()
            if text and text != full_name:
                print(f"Header title: replacing '{text}' -> {{{{CL_HEADLINE}}}}")
                replace_para_text(para, "{{CL_HEADLINE}}")
                break

    # Static paragraphs to leave unchanged
    STATIC_STARTS = {"Madame", "Monsieur", "Cordialement", full_name}

    placeholders = [
        "{{DATE}}",
        "{{COMPANY_ADDRESSEE}}",
        "{{COMPANY_LOCATION}}",
        "{{SUBJECT_LINE}}",
        "{{OPENING_PARA}}",
        "{{BODY_PARA_1}}",
        "{{BODY_PARA_2}}",
        "{{BODY_PARA_3}}",
        "{{CLOSING_PARA}}",
    ]

    ph_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(text.startswith(s) for s in STATIC_STARTS):
            continue
        if ph_index >= len(placeholders):
            break
        print(f"Replacing '{text[:60]}' -> {placeholders[ph_index]}")
        replace_para_text(para, placeholders[ph_index])
        ph_index += 1

    TEMPLATES_DIR.mkdir(exist_ok=True)
    output_path = TEMPLATES_DIR / output_name
    doc.save(str(output_path))
    print(f"\nTemplate saved: {output_path}")
    print(f"Replaced {ph_index} variable paragraphs.")


if __name__ == "__main__":
    main()
